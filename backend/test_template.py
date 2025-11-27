"""Test reading template file with image in header"""

from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
import io

SCOPES = ['https://www.googleapis.com/auth/drive']

def authenticate():
    """Load credentials from token.json"""
    creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    return build('drive', 'v3', credentials=creds)

def find_folder(service, folder_name):
    """Find folder by name"""
    query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'"
    results = service.files().list(q=query, fields='files(id, name)').execute()
    folders = results.get('files', [])
    return folders[0]['id'] if folders else None

def get_file(service, folder_id, file_name):
    """Get file info from folder"""
    query = f"'{folder_id}' in parents and name contains '{file_name}'"
    results = service.files().list(q=query, fields='files(id, name, mimeType)').execute()
    files = results.get('files', [])
    return files[0] if files else None

def download_docx(service, file_id, mime_type):
    """Download or export file as docx"""
    fh = io.BytesIO()

    if mime_type == 'application/vnd.google-apps.document':
        request = service.files().export_media(
            fileId=file_id,
            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        request = service.files().get_media(fileId=file_id)

    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()

    fh.seek(0)
    return fh

def analyze_template(file_handle):
    """Analyze template structure"""
    doc = Document(file_handle)

    print("\n=== TEMPLATE ANALYSIS ===\n")

    # Check headers
    print("Headers:")
    for section in doc.sections:
        header = section.header
        if header.paragraphs:
            for para in header.paragraphs:
                if para.text.strip():
                    print(f"  Text: {para.text}")
        # Check for images in header
        if hasattr(header, '_element'):
            print(f"  Has header element: Yes")

    # Main content
    print("\nMain content paragraphs:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  [{i}] {para.text[:100]}")

    # Tables
    if doc.tables:
        print(f"\nTables found: {len(doc.tables)}")
        for i, table in enumerate(doc.tables):
            print(f"  Table {i}: {len(table.rows)} rows, {len(table.columns)} cols")

if __name__ == '__main__':
    service = authenticate()
    print("✓ Authenticated")

    folder_id = find_folder(service, 'dan_wine')
    print("✓ Found dan_wine folder")

    file_info = get_file(service, folder_id, 'template')
    if not file_info:
        print("✗ template file not found")
        exit(1)

    print(f"✓ Found: {file_info['name']}")

    file_handle = download_docx(service, file_info['id'], file_info['mimeType'])
    print("✓ Downloaded")

    analyze_template(file_handle)
