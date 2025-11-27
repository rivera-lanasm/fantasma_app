"""Generate wine tasting sheet from template and product data"""

from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from docx import Document
from docx.shared import Pt
import pandas as pd
import io
import os
import logging

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)

SCOPES = ['https://www.googleapis.com/auth/drive']

def authenticate():
    """Load credentials from token.json"""
    creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    return build('drive', 'v3', credentials=creds)

def find_folder(service, folder_name):
    """Find folder by name"""
    query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'"
    results = service.files().list(q=query, fields='files(id, name)').execute()
    folders = results.get('files', [])
    return folders[0]['id'] if folders else None

def get_file(service, folder_id, file_name):
    """Get file info from folder"""
    query = f"'{folder_id}' in parents and name contains '{file_name}'"
    results = service.files().list(q=query, fields='files(id, name, mimeType)').execute()
    files = results.get('files', [])
    return files[0] if files else None

def download_file(service, file_id, mime_type):
    """Download or export file"""
    fh = io.BytesIO()

    if mime_type == 'application/vnd.google-apps.document':
        request = service.files().export_media(
            fileId=file_id,
            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        request = service.files().get_media(fileId=file_id)

    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()

    fh.seek(0)
    return fh

def replace_placeholders(text, row, wine_num):
    """Replace quoted placeholders with actual data"""
    # Template uses Unicode curly quotes (chr(8220) and chr(8221))
    lq = chr(8220)  # Left double quote
    rq = chr(8221)  # Right double quote

    original_text = text

    replacements = {
        f'{lq}PRODUCER{rq}': str(row['PRODUCER']),
        f'{lq}REGION_APPELLATION{rq}': str(row['REGION_APPELLATION']),
        f'{lq}CUVEE_NAME{rq}': str(row['CUVEE_NAME']),
        f'{lq}VINTAGE{rq}': str(int(row['VINTAGE'])) if pd.notna(row['VINTAGE']) and str(row['VINTAGE']).replace('.','').isdigit() else str(row['VINTAGE']) if pd.notna(row['VINTAGE']) else '',
        f'{lq}BLEND_DETAILS{rq}': str(row['BLEND_DETAILS']),
        f'{lq}PACKAGING{rq}': str(row['PACKAGING']),
        f'{lq}STANDARD_PRICE{rq}': f"${row['STANDARD_PRICE']}"
    }

    # Handle discount price and ** marker
    if pd.notna(row['DISCOUNT_PRICE']):
        replacements[f', {lq}DISCOUNTED_PRICE{rq}**'] = f", ${row['DISCOUNT_PRICE']}**"
    else:
        # Remove comma, discounted price placeholder, and **
        replacements[f', {lq}DISCOUNTED_PRICE{rq}**'] = ""

    for placeholder, value in replacements.items():
        if placeholder in text:
            text = text.replace(placeholder, value)

    # Log if this is a new wine (first line)
    if lq + 'PRODUCER' in original_text:
        logger.info(f"\n--- Wine #{wine_num}: {row['PRODUCER']} ---")
        logger.info(f"  Before: {original_text}")
        logger.info(f"  After:  {text}")

    return text

def copy_paragraph_with_formatting(template_para, doc, row, wine_num):
    """Copy paragraph preserving all formatting and replace placeholders"""
    # Create new paragraph with same style
    new_para = doc.add_paragraph()
    new_para.style = template_para.style
    new_para.paragraph_format.alignment = template_para.paragraph_format.alignment

    # Copy each run (formatted text segment) from template
    for run in template_para.runs:
        new_run = new_para.add_run()

        # Replace placeholders in the run's text
        new_run.text = replace_placeholders(run.text, row, wine_num)

        # Copy all formatting
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.color.rgb = run.font.color.rgb

    return new_para

def generate_document(template_handle, sheet_handle):
    """Generate filled document from template and data"""
    # Load template (this preserves all formatting and images)
    doc = Document(template_handle)
    logger.info(f"\n=== TEMPLATE LOADED ===")
    logger.info(f"Total paragraphs in template: {len(doc.paragraphs)}")

    # Load product data
    df = pd.read_excel(sheet_handle)
    logger.info(f"\n=== SHEET DATA ===")
    logger.info(f"Total rows in sheet: {len(df)}")

    # Filter to only rows with "x" in Chosen column
    if 'Chosen' in df.columns:
        df = df[df['Chosen'].astype(str).str.strip().str.lower() == 'x']
        logger.info(f"Rows with 'x' in Chosen column: {len(df)}")

    logger.info(f"\n=== PROCESSING {len(df)} WINES ===")

    # Store template paragraphs for wine block (paragraphs 1-4, indices 1-4)
    wine_template_paras = [doc.paragraphs[i] for i in range(1, 5)]

    logger.info(f"\nWine template block:")
    for i, para in enumerate(wine_template_paras, 1):
        logger.info(f"  Line {i}: {para.text}")

    # Store footer paragraph (paragraph 14)
    footer_para_template = doc.paragraphs[14]
    logger.info(f"\nFooter: {footer_para_template.text}")

    # Delete placeholder content paragraphs (keep header, delete body placeholders)
    # Delete in reverse to avoid index shifting
    for i in range(len(doc.paragraphs) - 1, 0, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)

    logger.info(f"\n=== GENERATING DOCUMENT ===")

    # Process each wine
    for wine_num, (idx, row) in enumerate(df.iterrows(), 1):
        # Add wine block with formatting preserved
        for template_para in wine_template_paras:
            copy_paragraph_with_formatting(template_para, doc, row, wine_num)

        # Add 2 blank lines between wines
        doc.add_paragraph()
        doc.add_paragraph()

    # Add footer with formatting preserved
    copy_paragraph_with_formatting(footer_para_template, doc, df.iloc[0], 0)

    logger.info(f"\n=== DOCUMENT COMPLETE ===")
    logger.info(f"Final paragraph count: {len(doc.paragraphs)}")
    logger.info(f"Expected: {len(df)} wines x 6 lines (4 wine + 2 blank) + 1 footer = {len(df) * 6 + 1} paragraphs")

    return doc

def upload_document(service, folder_id, file_path):
    """Upload document to Drive"""
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [folder_id]
    }
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    return file.get('id')

if __name__ == '__main__':
    service = authenticate()
    print("✓ Authenticated")

    folder_id = find_folder(service, 'dan_wine')
    print("✓ Found dan_wine folder")

    # Download template
    template_info = get_file(service, folder_id, 'template')
    template_handle = download_file(service, template_info['id'], template_info['mimeType'])
    print("✓ Downloaded template")

    # Download product data
    sheet_info = get_file(service, folder_id, 'product_data')
    sheet_handle = download_file(service, sheet_info['id'], sheet_info['mimeType'])
    print("✓ Downloaded product data")

    # Generate document
    doc = generate_document(template_handle, sheet_handle)
    print("✓ Generated document")

    # Save locally
    output_file = 'tasting_sheet_output.docx'
    doc.save(output_file)
    print(f"✓ Saved to {output_file}")

    # Upload to Drive
    file_id = upload_document(service, folder_id, output_file)
    print(f"✓ Uploaded to Drive (ID: {file_id})")
