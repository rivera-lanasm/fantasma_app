"""Verify formatting matches between template and output"""

from docx import Document
import os

def analyze_paragraph_formatting(para, para_num):
    """Analyze and display paragraph formatting details"""
    print(f"\n  Paragraph {para_num}: {para.text[:50]}...")
    print(f"    Style: {para.style.name}")
    print(f"    Alignment: {para.paragraph_format.alignment}")

    for i, run in enumerate(para.runs):
        if run.text.strip():
            print(f"    Run {i}: '{run.text[:30]}'")
            print(f"      Bold: {run.bold}")
            print(f"      Italic: {run.italic}")
            print(f"      Font name: {run.font.name}")
            print(f"      Font size: {run.font.size}")

if __name__ == '__main__':
    print("=== TEMPLATE FORMATTING ===")
    template = Document('template.docx') if os.path.exists('template.docx') else None

    if not template:
        # Download from Drive
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseDownload
        import io

        creds = Credentials.from_authorized_user_file('token.json')
        service = build('drive', 'v3', credentials=creds)

        # Find dan_wine folder
        query = "name='dan_wine' and mimeType='application/vnd.google-apps.folder'"
        folders = service.files().list(q=query, fields='files(id)').execute().get('files', [])
        folder_id = folders[0]['id']

        # Get template
        query = f"'{folder_id}' in parents and name contains 'template'"
        files = service.files().list(q=query, fields='files(id, mimeType)').execute().get('files', [])
        file_info = files[0]

        # Download
        fh = io.BytesIO()
        if file_info['mimeType'] == 'application/vnd.google-apps.document':
            request = service.files().export_media(fileId=file_info['id'], mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            request = service.files().get_media(fileId=file_info['id'])

        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        template = Document(fh)

    print("\nTemplate wine block paragraphs (1-4):")
    for i in range(1, 5):
        analyze_paragraph_formatting(template.paragraphs[i], i)

    print("\n\n=== OUTPUT FORMATTING ===")
    output = Document('tasting_sheet_output.docx')

    print("\nFirst wine block paragraphs (first 4 content paragraphs):")
    # Skip any header paragraphs, find first content
    content_para_count = 0
    for i, para in enumerate(output.paragraphs):
        if para.text.strip() and content_para_count < 4:
            content_para_count += 1
            analyze_paragraph_formatting(para, content_para_count)
