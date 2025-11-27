"""Generate wine tasting sheet for specific row indices"""

from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from docx import Document
from docx.shared import Pt
import pandas as pd
import io
import os
import logging
import sys
from datetime import datetime
import re

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)

SCOPES = ['https://www.googleapis.com/auth/drive']

def authenticate():
    """Load credentials from token.json"""
    creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    return build('drive', 'v3', credentials=creds)

def find_folder(service, folder_name):
    """Find folder by name"""
    query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'"
    results = service.files().list(q=query, fields='files(id, name)').execute()
    folders = results.get('files', [])
    return folders[0]['id'] if folders else None

def get_file(service, folder_id, file_name):
    """Get file info from folder"""
    query = f"'{folder_id}' in parents and name contains '{file_name}'"
    results = service.files().list(q=query, fields='files(id, name, mimeType)').execute()
    files = results.get('files', [])
    return files[0] if files else None

def get_latest_xlsx(service, folder_id):
    """Get the most recently modified .xlsx file in folder"""
    query = f"'{folder_id}' in parents and mimeType='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'"
    results = service.files().list(
        q=query,
        fields='files(id, name, mimeType, modifiedTime)',
        orderBy='modifiedTime desc'
    ).execute()
    files = results.get('files', [])
    return files[0] if files else None

def download_file(service, file_id, mime_type):
    """Download or export file"""
    fh = io.BytesIO()

    if mime_type == 'application/vnd.google-apps.document':
        request = service.files().export_media(
            fileId=file_id,
            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        request = service.files().get_media(fileId=file_id)

    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()

    fh.seek(0)
    return fh

def replace_placeholders(text, row, wine_num):
    """Replace quoted placeholders with actual data"""
    lq = chr(8220)
    rq = chr(8221)

    original_text = text

    replacements = {
        f'{lq}PRODUCER{rq}': str(row['PRODUCER']),
        f'{lq}REGION_APPELLATION{rq}': str(row['REGION_APPELLATION']),
        f'{lq}CUVEE_NAME{rq}': str(row['CUVEE_NAME']),
        f'{lq}VINTAGE{rq}': str(int(row['VINTAGE'])) if pd.notna(row['VINTAGE']) and str(row['VINTAGE']).replace('.','').isdigit() else str(row['VINTAGE']) if pd.notna(row['VINTAGE']) else '',
        f'{lq}BLEND_DETAILS{rq}': str(row['BLEND_DETAILS']),
        f'{lq}PACKAGING{rq}': str(row['PACKAGING']),
        f'{lq}STANDARD_PRICE{rq}': f"${row['STANDARD_PRICE']}"
    }

    if pd.notna(row['DISCOUNT_PRICE']):
        replacements[f', {lq}DISCOUNTED_PRICE{rq}**'] = f", ${row['DISCOUNT_PRICE']}**"
    else:
        replacements[f', {lq}DISCOUNTED_PRICE{rq}**'] = ""

    for placeholder, value in replacements.items():
        if placeholder in text:
            text = text.replace(placeholder, value)

    if lq + 'PRODUCER' in original_text:
        logger.info(f"\n--- Wine #{wine_num}: {row['PRODUCER']} ---")
        logger.info(f"  Before: {original_text}")
        logger.info(f"  After:  {text}")

    return text

def copy_paragraph_with_formatting(template_para, doc, row, wine_num):
    """Copy paragraph preserving all formatting and replace placeholders"""
    new_para = doc.add_paragraph()
    new_para.style = template_para.style
    new_para.paragraph_format.alignment = template_para.paragraph_format.alignment

    for run in template_para.runs:
        new_run = new_para.add_run()
        new_run.text = replace_placeholders(run.text, row, wine_num)

        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.color.rgb = run.font.color.rgb

    return new_para

def generate_document(template_handle, sheet_handle, row_indices):
    """Generate filled document from template and specific rows"""
    doc = Document(template_handle)
    logger.info(f"\n=== TEMPLATE LOADED ===")

    df = pd.read_excel(sheet_handle)
    logger.info(f"\n=== SHEET DATA ===")
    logger.info(f"Total rows in sheet: {len(df)}")

    # Filter to specified row indices
    df = df.iloc[row_indices]
    logger.info(f"Selected {len(df)} rows: {row_indices}")

    logger.info(f"\n=== PROCESSING {len(df)} WINES ===")

    # Get template paragraphs (handle templates with different lengths)
    if len(doc.paragraphs) < 5:
        wine_template_paras = doc.paragraphs[1:min(5, len(doc.paragraphs))]
    else:
        wine_template_paras = [doc.paragraphs[i] for i in range(1, 5)]

    footer_idx = min(14, len(doc.paragraphs) - 1) if len(doc.paragraphs) > 14 else len(doc.paragraphs) - 1
    footer_para_template = doc.paragraphs[footer_idx] if len(doc.paragraphs) > 0 else None

    # Delete placeholder content
    for i in range(len(doc.paragraphs) - 1, 0, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)

    # Process each wine
    for wine_num, (idx, row) in enumerate(df.iterrows(), 1):
        for template_para in wine_template_paras:
            copy_paragraph_with_formatting(template_para, doc, row, wine_num)

        doc.add_paragraph()
        doc.add_paragraph()

    # Add footer if template has one
    if footer_para_template:
        copy_paragraph_with_formatting(footer_para_template, doc, df.iloc[0], 0)

    logger.info(f"\n=== DOCUMENT COMPLETE ===")
    return doc

def generate_price_list(template_handle, df):
    """Generate price list from template and data"""
    doc = Document(template_handle)
    logger.info(f"\n=== PRICE LIST TEMPLATE LOADED ===")

    if not doc.tables:
        logger.info("No tables found in price list template")
        return doc

    table = doc.tables[0]
    logger.info(f"Table has {len(table.rows)} rows, {len(table.columns)} columns")

    # Remove template rows (keep header row 0, delete placeholder rows)
    for i in range(len(table.rows) - 1, 0, -1):
        table._element.remove(table.rows[i]._element)

    # Add a row for each wine
    for idx, row in df.iterrows():
        new_row = table.add_row()

        # Column 0: Producer, Cuvee, Vintage, Blend, Region
        producer_cell = new_row.cells[0]
        producer_text = f"{row['PRODUCER']}\n"
        producer_text += f"{row['CUVEE_NAME']}"
        if pd.notna(row['VINTAGE']):
            vintage_str = str(int(row['VINTAGE'])) if str(row['VINTAGE']).replace('.','').isdigit() else str(row['VINTAGE'])
            producer_text += f" {vintage_str}"
        producer_text += f"\n({row['BLEND_DETAILS']})\n"
        producer_text += f"{row['REGION_APPELLATION']}"
        producer_cell.text = producer_text

        # Column 1: Standard Price
        new_row.cells[1].text = f"${row['STANDARD_PRICE']}"

        # Column 2: Discount Price
        if pd.notna(row['DISCOUNT_PRICE']):
            new_row.cells[2].text = f"${row['DISCOUNT_PRICE']}"
        else:
            new_row.cells[2].text = ""

    logger.info(f"Added {len(df)} wines to price list")
    return doc

def get_timestamped_filename(service, folder_id, base_name, extension='.docx'):
    """Generate timestamped filename with version number"""
    today = datetime.now().strftime('%Y-%m-%d')

    # List existing files in folder with same date and base name
    query = f"'{folder_id}' in parents and name contains '{today}' and name contains '{base_name}'"
    results = service.files().list(q=query, fields='files(name)').execute()
    existing_files = results.get('files', [])

    # Find highest version number
    version = 1
    for file in existing_files:
        # Match pattern: base_name_YYYY-MM-DD_vN.docx
        match = re.search(rf'{today}_v(\d+)', file['name'])
        if match:
            file_version = int(match.group(1))
            if file_version >= version:
                version = file_version + 1

    return f"{base_name}_{today}_v{version}{extension}"

def upload_document(service, folder_id, file_path):
    """Upload document to Drive"""
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [folder_id]
    }
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    return file.get('id')

if __name__ == '__main__':
    # Get row indices from command line args
    if len(sys.argv) < 2:
        print("Usage: python generate_selected_wines.py ROW1 ROW2 ...")
        sys.exit(1)

    row_indices = [int(arg) for arg in sys.argv[1:]]

    service = authenticate()
    print("✓ Authenticated")

    folder_id = find_folder(service, 'Automation Demo Folder')
    print("✓ Found Automation Demo Folder")

    # Download tasting sheet template
    template_info = get_file(service, folder_id, 'TASTING SHEET')
    template_handle = download_file(service, template_info['id'], template_info['mimeType'])
    print("✓ Downloaded tasting sheet template")

    # Download price list template
    price_template_info = get_file(service, folder_id, 'Price list')
    price_template_handle = download_file(service, price_template_info['id'], price_template_info['mimeType'])
    print("✓ Downloaded price list template")

    # Download product data
    sheet_info = get_latest_xlsx(service, folder_id)
    if not sheet_info:
        print("✗ No Excel files found in folder")
        sys.exit(1)
    print(f"✓ Found latest Excel: {sheet_info['name']}")
    sheet_handle = download_file(service, sheet_info['id'], sheet_info['mimeType'])
    print("✓ Downloaded product data")

    # Read full dataframe
    df_full = pd.read_excel(sheet_handle)
    df_selected = df_full.iloc[row_indices]

    # Generate tasting sheet
    sheet_handle.seek(0)  # Reset for reading again
    tasting_doc = generate_document(template_handle, sheet_handle, row_indices)
    print("✓ Generated tasting sheet")

    # Generate price list
    price_doc = generate_price_list(price_template_handle, df_selected)
    print("✓ Generated price list")

    # Save with timestamped filenames
    tasting_filename = get_timestamped_filename(service, folder_id, 'Tasting_Sheet')
    price_filename = get_timestamped_filename(service, folder_id, 'Price_List')

    tasting_doc.save(tasting_filename)
    print(f"✓ Saved tasting sheet: {tasting_filename}")

    price_doc.save(price_filename)
    print(f"✓ Saved price list: {price_filename}")

    # Upload both to Drive
    tasting_id = upload_document(service, folder_id, tasting_filename)
    print(f"✓ Uploaded tasting sheet to Drive")

    price_id = upload_document(service, folder_id, price_filename)
    print(f"✓ Uploaded price list to Drive")
